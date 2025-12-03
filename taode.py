import streamlit as st
import json
import random
import io
import zipfile
import math
from copy import deepcopy

# --- THƯ VIỆN PDF ---
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

# --- THƯ VIỆN WORD ---
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ==========================================
# 1. CẤU HÌNH FONT & HELPER
# ==========================================

def register_fonts():
    """Đăng ký font Times New Roman cho PDF"""
    font_regular = 'Times.ttf'
    font_bold = 'Timesbd.ttf' # Times New Roman Bold
    font_italic = 'Timesi.ttf'
    
    used_font = 'Helvetica' # Fallback
    
    try:
        # Ưu tiên load Times New Roman
        pdfmetrics.registerFont(TTFont('Times-Roman', font_regular))
        used_font = 'Times-Roman'
        
        # Thử load thêm font đậm (nếu có)
        try:
            pdfmetrics.registerFont(TTFont('Times-Bold', font_bold))
        except:
            # Nếu không có file đậm, map font đậm về font thường (không khuyến khích)
            pass
            
    except:
        pass # Dùng mặc định nếu không tìm thấy file
        
    return used_font

def format_text_pdf(text):
    if not text: return ""
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    return text.replace('\n', '<br/>').replace('  ', '&nbsp;&nbsp;')

# ==========================================
# 2. XỬ LÝ PDF BOOKMARK (Heading cho PDF)
# ==========================================

class BookmarkCanvas(canvas.Canvas):
    """
    Canvas tùy chỉnh để tạo Bookmark (Heading) tự động trong PDF
    """
    def __init__(self, *args, **kwargs):
        canvas.Canvas.__init__(self, *args, **kwargs)
        self._bookmarks = []

    def add_bookmark(self, name, key):
        self._bookmarks.append((name, key))
    
    # Hàm này sẽ được gọi mỗi khi flowable được vẽ (cần logic phức tạp hơn để gắn đúng vị trí)
    # Tuy nhiên, để đơn giản cho Streamlit, ta sẽ dùng phương pháp addOutlineEntry trực tiếp trong loop

# ==========================================
# 3. LOGIC TRỘN (GIỮ NGUYÊN)
# ==========================================

def mix_exam_data(original_questions, shuffle_questions=True, shuffle_options=True, balance_distribution=True):
    questions = deepcopy(original_questions)
    
    if shuffle_questions:
        random.shuffle(questions)
        for idx, q in enumerate(questions):
            q['display_id'] = idx + 1 
    else:
        for q in questions:
            q['display_id'] = q['id']

    if shuffle_options:
        total_q = len(questions)
        target_indices = []
        if balance_distribution:
            base = [0, 1, 2, 3]
            repeats = (total_q // 4) + 1
            pool = (base * repeats)[:total_q]
            random.shuffle(pool)
            target_indices = pool
        
        for idx, q in enumerate(questions):
            opts = q.get('options', [])
            correct_idx = q.get('correct_answer_index', -1)
            
            if opts and correct_idx != -1 and len(opts) == 4:
                correct_text = opts[correct_idx]
                distractors = [o for i, o in enumerate(opts) if i != correct_idx]
                random.shuffle(distractors)
                
                new_correct_idx = target_indices[idx] if balance_distribution else random.randint(0, 3)
                new_opts = [None] * 4
                new_opts[new_correct_idx] = correct_text
                d_ptr = 0
                for i in range(4):
                    if new_opts[i] is None:
                        new_opts[i] = distractors[d_ptr]
                        d_ptr += 1
                q['options'] = new_opts
                q['correct_answer_index'] = new_correct_idx
            elif shuffle_options: 
                paired = list(zip(opts, [i==correct_idx for i in range(len(opts))]))
                random.shuffle(paired)
                q['options'] = [p[0] for p in paired]
                for i, p in enumerate(paired):
                    if p[1]: q['correct_answer_index'] = i; break
    return questions

# ==========================================
# 4. TẠO PDF (CHUẨN VIỆT NAM + BOOKMARKS)
# ==========================================

def generate_pdf_bytes(questions, mode="exam"):
    buffer = io.BytesIO()
    
    # 1. Cấu hình trang chuẩn: Lề trái 3cm, Phải/Trên/Dưới 2cm
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                            leftMargin=3*cm, rightMargin=2*cm, 
                            topMargin=2*cm, bottomMargin=2*cm)
    
    story = []
    font_name = register_fonts()
    font_bold_name = 'Times-Bold' if font_name == 'Times-Roman' else font_name # Fallback
    
    # Styles chuẩn
    styles = getSampleStyleSheet()
    
    # Header Style
    style_header_school = ParagraphStyle('HSchool', fontName=font_bold_name, fontSize=11, alignment=TA_CENTER)
    style_header_exam = ParagraphStyle('HExam', fontName=font_bold_name, fontSize=12, alignment=TA_CENTER)
    
    # Question Style (Size 13pt chuẩn)
    style_q = ParagraphStyle('Quest', parent=styles['Normal'], fontName=font_name, fontSize=13, leading=16, spaceAfter=6, alignment=TA_JUSTIFY)
    style_opt = ParagraphStyle('Opt', parent=styles['Normal'], fontName=font_name, fontSize=13, leading=16)

    # --- TẠO HEADER ---
    # Bảng Header 2 cột: Sở/Trường bên trái, Tên thi bên phải
    h_text_left = "SỞ GD&ĐT ........................<br/>TRƯỜNG THPT ........................"
    h_text_right = f"<b>{'ĐỀ THI TRẮC NGHIỆM' if mode == 'exam' else 'ĐÁP ÁN'}</b><br/>Môn: Tin học"
    
    h_table = Table([[Paragraph(h_text_left, style_header_school), Paragraph(h_text_right, style_header_exam)]], 
                    colWidths=[8*cm, 8*cm])
    h_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
    ]))
    story.append(h_table)
    story.append(Spacer(1, 0.5*cm))
    # Kẻ đường ngang
    story.append(Paragraph("_______________________________________", ParagraphStyle('Line', alignment=TA_CENTER)))
    story.append(Spacer(1, 1*cm))

    if mode == "exam":
        labels = ["A.", "B.", "C.", "D."]
        
        # Để tạo Bookmark, ta cần custom Canvas loop. 
        # Cách đơn giản nhất trong SimpleDocTemplate là chèn Anchor.
        
        for q in questions:
            # Tạo Anchor cho Bookmark
            key = f"cau_{q['display_id']}"
            q_text_content = format_text_pdf(q['question'])
            
            # Thẻ <a> ẩn để làm điểm neo bookmark (nâng cao)
            # Hoặc đơn giản là in câu hỏi
            full_q_text = f"<b>Câu {q['display_id']}:</b> {q_text_content}"
            
            # Đoạn văn câu hỏi
            p = Paragraph(full_q_text, style_q)
            story.append(p)
            
            # -- ĐÁP ÁN (LAYOUT A-C / B-D) --
            opts = q.get('options', [])
            clean_opts = [str(o) for o in opts]
            opt_paras = []
            for i, o_text in enumerate(clean_opts):
                if i < 4:
                    opt_paras.append(Paragraph(f"<b>{labels[i]}</b> {format_text_pdf(o_text)}", style_opt))

            max_len = max([len(o) for o in clean_opts]) if clean_opts else 0
            table_data = []
            col_widths = []

            # Logic chia cột
            doc_width = A4[0] - 5*cm # Trừ lề trái 3 + phải 2
            
            if len(opt_paras) == 4 and max_len < 40:
                row1 = [opt_paras[0], opt_paras[2]]
                row2 = [opt_paras[1], opt_paras[3]]
                table_data = [row1, row2]
                col_widths = [doc_width/2, doc_width/2]
            else:
                table_data = [[o] for o in opt_paras]
                col_widths = [doc_width]

            if table_data:
                t = Table(table_data, colWidths=col_widths)
                t.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                ]))
                story.append(t)
            story.append(Spacer(1, 0.3*cm))

    else: # KEY
        data = []; row = []
        lbls = ["A", "B", "C", "D"]
        for q in questions:
            idx = q.get('correct_answer_index', -1)
            ans = lbls[idx] if idx != -1 else "?"
            row.append(f"{q['display_id']}: {ans}")
            if len(row) == 5: data.append(row); row = []
        if row: 
            while len(row) < 5: row.append("")
            data.append(row)
        
        if data:
            t = Table(data, colWidths=[2.5*cm]*5, rowHeights=0.8*cm)
            t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('FONTNAME', (0,0), (-1,-1), font_name)]))
            story.append(t)

    # --- HÀM BUILD ĐẶC BIỆT ĐỂ TẠO BOOKMARK ---
    def add_bookmarks(canvas, doc):
        """Hàm callback để vẽ bookmark lên từng trang"""
        canvas.saveState()
        # Vì SimpleDocTemplate khó biết chính xác vị trí Y của từng câu trong flow
        # Nên ở đây ta tạo bookmark giả lập theo trang hoặc thêm logic phức tạp hơn.
        # Để đơn giản và hiệu quả: Ta chỉ bookmark trang đầu hoặc các mục lớn.
        # Mở rộng: Muốn bookmark từng câu chính xác cần dùng Paragraph(..., destination='...')
        canvas.restoreState()

    # Lưu ý: Với SimpleDocTemplate, việc tạo Bookmark trỏ đúng từng câu khá phức tạp
    # Dưới đây là cách hack để tạo bookmark mỗi khi gặp một Flowable nhất định (nâng cao)
    # Trong phạm vi code này, ta sẽ build bình thường, font và layout đã chuẩn.
    
    doc.build(story)
    buffer.seek(0)
    return buffer

# ==========================================
# 5. TẠO WORD (CHUẨN VIỆT NAM + HEADING)
# ==========================================
def fix_indent_word(text):
    """
    Chuyển đổi các khoảng trắng đầu dòng thành Non-breaking space (\u00A0)
    để Word bắt buộc phải hiển thị, không được tự động co lại.
    """
    if not text: return ""
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        # Tách phần nội dung và phần khoảng trắng đầu dòng
        stripped_content = line.lstrip()
        num_spaces = len(line) - len(stripped_content)
        
        if num_spaces > 0:
            # Thay thế bằng \u00A0 (Non-breaking space)
            # Ký tự này trong Word có độ rộng cố định và không bị xóa
            indent_str = "\u00A0" * num_spaces
            processed_lines.append(indent_str + stripped_content)
        else:
            processed_lines.append(line)
            
    return '\n'.join(processed_lines)

# ==========================================
# 4. TẠO WORD (ĐÃ SỬA LỖI MẤT KHOẢNG TRẮNG)
# ==========================================

def generate_word_bytes(questions, mode="exam"):
    buffer = io.BytesIO()
    doc = Document()
    
    # 1. Cấu hình trang chuẩn
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    def set_font(run, bold=False, size=13):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # 2. Header
    h_table = doc.add_table(rows=1, cols=2)
    h_table.autofit = False
    h_table.columns[0].width = Cm(9)
    h_table.columns[1].width = Cm(7)
    
    c1 = h_table.cell(0, 0)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("SỞ GD&ĐT ........................\nTRƯỜNG THPT ........................")
    set_font(r1, bold=True, size=11)
    
    c2 = h_table.cell(0, 1)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_str = "ĐỀ THI TRẮC NGHIỆM" if mode == "exam" else "ĐÁP ÁN"
    r2 = p2.add_run(f"{title_str}\nMÔN: TIN HỌC")
    set_font(r2, bold=True, size=12)

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_line = p_line.add_run("____________________________________")
    set_font(r_line, bold=True)
    doc.add_paragraph() 

    if mode == "exam":
        labels = ["A.", "B.", "C.", "D."]
        for q in questions:
            # --- HEADING CÂU HỎI ---
            heading = doc.add_heading(level=2)
            heading.paragraph_format.space_before = Pt(6)
            heading.paragraph_format.space_after = Pt(6)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            r_num = heading.add_run(f"Câu {q['display_id']}: ")
            set_font(r_num, bold=True, size=13)
            
            # --- QUAN TRỌNG: GỌI HÀM SỬA LỖI THỤT DÒNG ---
            # Xử lý text câu hỏi trước khi đưa vào Word
            clean_question_text = fix_indent_word(q['question'])
            
            r_content = heading.add_run(clean_question_text)
            set_font(r_content, bold=False, size=13)

            # --- ĐÁP ÁN ---
            opts = q.get('options', [])
            clean_opts = [str(o) for o in opts]
            max_len = max([len(o) for o in clean_opts]) if clean_opts else 0

            if len(clean_opts) == 4 and max_len < 40:
                table = doc.add_table(rows=2, cols=2)
                table.autofit = True
                map_pos = [(0,0), (1,0), (0,1), (1,1)]
                for i in range(4):
                    r_idx, c_idx = map_pos[i]
                    cell = table.cell(r_idx, c_idx)
                    p_opt = cell.paragraphs[0]
                    # Cũng sửa thụt dòng cho đáp án (phòng hờ)
                    fixed_opt = fix_indent_word(clean_opts[i])
                    run_opt = p_opt.add_run(f"{labels[i]} {fixed_opt}")
                    set_font(run_opt, size=13)
            else:
                table = doc.add_table(rows=len(clean_opts), cols=1)
                for i, txt in enumerate(clean_opts):
                    cell = table.cell(i, 0)
                    p_opt = cell.paragraphs[0]
                    fixed_opt = fix_indent_word(txt)
                    run_opt = p_opt.add_run(f"{labels[i]} {fixed_opt}")
                    set_font(run_opt, size=13)
            
            doc.add_paragraph() 

    else: # KEY (Giữ nguyên)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        curr_row = table.rows[0]
        c_count = 0
        lbls = ["A", "B", "C", "D"]
        for q in questions:
            idx = q.get('correct_answer_index', -1)
            ans = lbls[idx] if idx != -1 else "?"
            cell = curr_row.cells[c_count]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{q['display_id']}: {ans}")
            set_font(run, bold=True, size=13)
            c_count += 1
            if c_count >= 5: c_count = 0; curr_row = table.add_row()

    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 6. UI STREAMLIT
# ==========================================

st.set_page_config(page_title="Exam Pro VN Standard", layout="wide", page_icon="🇻🇳")
st.title("🇻🇳 Exam Mixer Pro - Chuẩn Văn Bản Việt Nam")
st.markdown("""
**Tiêu chuẩn áp dụng:**
* Font: **Times New Roman** (Cần file .ttf cùng thư mục)
* Cỡ chữ: **13pt**
* Khổ giấy A4, Lề: **Trái 3cm**, Phải/Trên/Dưới **2cm**.
* Word: Có **Heading** (Navigation Pane).
""")

with st.sidebar:
    st.header("Upload")
    uploaded_file = st.file_uploader("Chọn file JSON", type=["json"])
    st.header("Cấu hình")
    shuffle_q = st.checkbox("Trộn câu hỏi", value=True)
    shuffle_o = st.checkbox("Trộn đáp án", value=True)
    balance_dist = st.checkbox("Cân bằng đáp án", value=True)

if uploaded_file:
    try:
        raw_data = json.load(uploaded_file)
        trigger_id = f"{uploaded_file.name}_{shuffle_q}_{shuffle_o}_{balance_dist}"
        
        if 'last_trigger' not in st.session_state or st.session_state.last_trigger != trigger_id:
            with st.spinner("Đang trộn đề..."):
                st.session_state.mixed_data = mix_exam_data(raw_data, shuffle_q, shuffle_o, balance_dist)
                st.session_state.last_trigger = trigger_id
        
        mixed_data = st.session_state.mixed_data
        
        # Thống kê
        st.divider()
        cnt = {"A":0, "B":0, "C":0, "D":0}
        lbls = ["A", "B", "C", "D"]
        for q in mixed_data:
            if q['correct_answer_index'] != -1: cnt[lbls[q['correct_answer_index']]] += 1
        cols = st.columns(4)
        for i, (k, v) in enumerate(cnt.items()): cols[i].metric(f"Đáp án {k}", f"{v}")

        # Download
        st.subheader("Tải về")
        c1, c2 = st.columns(2)
        with c1:
            pdf_exam = generate_pdf_bytes(mixed_data, "exam")
            pdf_key = generate_pdf_bytes(mixed_data, "key")
            zip_pdf = io.BytesIO()
            with zipfile.ZipFile(zip_pdf, "a", zipfile.ZIP_DEFLATED, False) as zf:
                zf.writestr("De_Thi_ChuanVN.pdf", pdf_exam.getvalue())
                zf.writestr("Dap_An.pdf", pdf_key.getvalue())
            st.download_button("📥 Tải PDF Chuẩn VN (.zip)", zip_pdf.getvalue(), "PDF_VN_Standard.zip", "application/zip", use_container_width=True)
            
        with c2:
            word_exam = generate_word_bytes(mixed_data, "exam")
            word_key = generate_word_bytes(mixed_data, "key")
            zip_word = io.BytesIO()
            with zipfile.ZipFile(zip_word, "a", zipfile.ZIP_DEFLATED, False) as zf:
                zf.writestr("De_Thi_ChuanVN.docx", word_exam.getvalue())
                zf.writestr("Dap_An.docx", word_key.getvalue())
            st.download_button("📥 Tải Word Chuẩn VN (.zip)", zip_word.getvalue(), "Word_VN_Standard.zip", "application/zip", use_container_width=True, type="primary")

    except Exception as e:
        st.error(f"Lỗi: {e}")